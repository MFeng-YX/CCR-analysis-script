#!/usr/bin/env python3
"""
pip install python-docx openpyxl
用法：
    python extract_column.py
然后按提示依次输入：
    1) Excel 完整路径
    2) 列字母，如 A/B/C...
    3) 起始行号
    4) 结束行号
    5) 输出格式：word / txt
"""

from pathlib import Path
import openpyxl
from docx import Document
from docx.shared import Pt

def main():
    # 1. 读用户输入
    excel_path = input("请输入 Excel 文件完整路径：").strip().strip('"')
    col_letter = input("请输入列字母（如 A、B）：").strip().upper()
    start_row = int(input("请输入起始行号：").strip())
    end_row   = int(input("请输入结束行号：").strip())
    out_type  = input("输出格式(word / txt)：").strip().lower()

    # 2. 校验
    excel_path = Path(excel_path).expanduser().resolve()
    if not excel_path.exists():
        print("文件不存在！")
        return
    if out_type not in {"word", "txt"}:
        print("输出格式只能填 word 或 txt")
        return

    # 3. 读数据
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    data_list = []
    for r in range(start_row, end_row + 1):
        val = ws[f"{col_letter}{r}"].value
        data_list.append((r, str(val) if val is not None else ""))

    # 4. 写文件
    suffix = ".docx" if out_type == "word" else ".txt"
    out_path = excel_path.with_suffix(suffix)

    if out_type == "word":
        doc = Document()
        for row_idx, content in data_list:
            p = doc.add_paragraph()
            run = p.add_run(f"{row_idx}：")
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(content)
            doc.add_paragraph()   # 额外空一行
        doc.save(out_path)
    else:  # txt
        with open(out_path, "w", encoding="utf-8") as f:
            for row_idx, content in data_list:
                f.write(f"*{row_idx}*：\n")
                f.write(f"{content}\n\n")
    print(f"已生成：{out_path}")

if __name__ == "__main__":
    main()